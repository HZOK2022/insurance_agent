"""摄取知识文件到 Qdrant(insurance_knowledge)。

用法: python scripts/ingest_kb.py --path <文件|目录> [--clear] [--limit N]
支持 .txt/.md(直接文本)、.pdf(pdfplumber 抽文本)、
.docx(python-docx 抽段落+表格)、.xlsx(openpyxl 抽工作表行列)。
"""
from __future__ import annotations
import argparse
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.config import load
from app.retrieval.chunker import chunk_documents
from app.retrieval.embedder import Embedder
from app.retrieval.qdrant_store import QdrantStore
from app.retrieval.knowledge_store import KnowledgeStore
from app.retrieval.categories import classify_product_category

_EXTS = (".txt", ".md", ".pdf", ".docx", ".xlsx")

# 扩展名 -> 入库 doc_type,便于下游按类型过滤/检索
_DOC_TYPE = {
    ".txt": "text",
    ".md": "markdown",
    ".pdf": "policy_pdf",
    ".docx": "policy_docx",
    ".xlsx": "rate_table",
}


def read_text(path: str):
    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt", ".md"):
        with open(path, encoding="utf-8") as f:
            return f.read()
    if ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join((p.extract_text() or "") for p in pdf.pages)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".xlsx":
        return _read_xlsx(path)
    return None


def _read_docx(path: str):
    """python-docx:正文段落 + 表格(逐行用 | 拼接)。"""
    from docx import Document
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for ti, table in enumerate(doc.tables, 1):
        parts.append(f"[表格{ti}]")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_xlsx(path: str):
    """openpyxl:逐工作表,表头行 + 数据行(逐行用 | 拼接)。"""
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"[{ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def build_docs(path: str, category: str = ""):
    text = read_text(path)
    if not text:
        return []
    base = os.path.splitext(os.path.basename(path))[0]
    ext = os.path.splitext(path)[1].lower()
    # 保险类别:优先显式 --category,否则按 doc_id 关键词判定(医疗险/重疾险/意外险/…)
    cat = category or classify_product_category(base, base)
    return [{"text": text,
             "meta": {"chunk_id": base, "doc_id": base, "version": "v1", "section": "",
                      "doc_type": _DOC_TYPE.get(ext, "policy_document"), "source": path,
                      "title": base, "product_category": cat}}]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--path", required=True)
    ap.add_argument("--clear", action="store_true")
    ap.add_argument("--limit", type=int, default=0, help="仅前 N 个文件(测试用)")
    ap.add_argument("--category", default="", help="保险类别(医疗险/重疾险/意外险/…);缺省按 doc_id 关键词判定")
    a = ap.parse_args()

    cfg = load()
    print(f"[cfg] collection={cfg.qdrant_collection} embedding={cfg.embedding_model} chunk={cfg.chunk_size}/{cfg.chunk_overlap}")

    # 事实源(SQLite chunks):canon,供重建 Qdrant / 版本化;Qdrant = 派生向量索引
    kstore = KnowledgeStore(getattr(cfg, "knowledge_db_path", "data/knowledge.db"))
    # 先删(若 --clear) —— Qdrant 集合 + SQLite 事实源都清,重建
    if a.clear:
        from qdrant_client import QdrantClient
        QdrantClient(url=cfg.qdrant_url).delete_collection(cfg.qdrant_collection)
        kstore.conn.execute("DELETE FROM chunks"); kstore.conn.commit()
        print("[clear] 已删除集合 + SQLite chunks 事实源")
    store = QdrantStore(cfg.qdrant_url, cfg.qdrant_collection, cfg.embedding_dim)

    files = []
    if os.path.isfile(a.path):
        files = [a.path]
    else:
        for root, _, fs in os.walk(a.path):
            for fn in fs:
                if fn.lower().endswith(_EXTS):
                    files.append(os.path.join(root, fn))
    if a.limit:
        files = files[:a.limit]

    docs = []
    for f in files:
        docs += build_docs(f, a.category)
    chunks = chunk_documents(docs, cfg.chunk_size, cfg.chunk_overlap)
    print(f"[files] {len(files)} -> [chunks] {len(chunks)}")

    # 先把 chunks 落 SQLite 事实源(唯一真相);再嵌入进 Qdrant(派生索引)
    kstore.upsert_chunks(chunks)
    print(f"[fact] chunks -> knowledge.db({kstore.count()})")
    embedder = Embedder(cfg.embedding_model, cfg.embedding_device, cfg.embedding_batch_size)
    B = 64
    i = 0
    while i < len(chunks):
        batch = chunks[i:i + B]
        vecs = embedder.embed([c["content"] for c in batch])
        store.upsert([{"vector": v, "content": c["content"], "meta": c["meta"]}
                      for c, v in zip(batch, vecs)])
        i += B
    print(f"[done] ingested {len(chunks)} chunks -> '{cfg.qdrant_collection}'")


if __name__ == "__main__":
    main()
